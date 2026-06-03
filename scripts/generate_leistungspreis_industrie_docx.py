"""Erzeugt die Industrie-Uebersichts-.docx (reale Netzentgelte + Buckets).

Gegliedert nach Normalfall (Jahresleistungspreis) und atypischem Fall
(atypische Netznutzung), jeweils mit Konzept, Beispiel, Solver-Bedeutung
und code-technischer Umsetzung. Mit eingebetteten Grafiken (matplotlib),
Titel-Banner, Ueberschriften-Linien und Seitenzahlen.

Reale Tarife: Bayernwerk Netz GmbH und Regensburg Netz GmbH (REWAG), 2026.
"""

import tempfile
from pathlib import Path

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import Patch

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


# ---- Farbpalette ----
BLUE = RGBColor(0x0B, 0x3D, 0x91)
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x55, 0x55, 0x55)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT = RGBColor(0xD3, 0xDE, 0xF2)

HEX_BLUE = "0B3D91"
HEX_ZEBRA = "F2F6FB"

C_PRIMARY = "#0B3D91"
C_RED = "#C0392B"
C_GREEN = "#1E8449"
C_GREY = "#7F8C8D"
C_LIGHTRED = "#F2D7D5"
C_LIGHTGREEN = "#D5F0DD"


# ===========================================================================
# Word-Hilfsfunktionen
# ===========================================================================

def _set_cell_bg(cell, hex_color: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
    })
    tcpr.append(shd)


def _add_bottom_border(paragraph, color=HEX_BLUE, size=10) -> None:
    pPr = paragraph._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    pbdr.append(bottom)
    pPr.append(pbdr)


def add_banner(doc, title, subtitle):
    """Farbiges Titel-Banner (einzelliges, gefuelltes Tabellenfeld)."""
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    _set_cell_bg(cell, HEX_BLUE)
    cell.width = Inches(6.3)
    cell.margins = None  # default
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(title)
    r.bold = True
    r.font.size = Pt(22)
    r.font.color.rgb = WHITE
    r.font.name = "Calibri"
    p2 = cell.add_paragraph()
    r2 = p2.add_run(subtitle)
    r2.font.size = Pt(10.5)
    r2.font.color.rgb = LIGHT
    r2.font.name = "Calibri"
    # etwas Innenabstand
    for par in cell.paragraphs:
        par.paragraph_format.space_before = Pt(2)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)


def add_heading(doc, text, size=15, color=BLUE, border=True,
                space_before=14, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    if border:
        _add_bottom_border(p)
    return p


def add_sub(doc, text, color=BLUE, size=12):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_body(doc, text, size=10.5, italic=False, color=DARK, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK
    return p


def add_mono(doc, lines, size=9.0):
    for ln in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Pt(12)
        run = p.add_run(ln)
        run.font.size = Pt(size)
        run.font.name = "Consolas"
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(8)


def add_figure(doc, path, width_in, caption):
    doc.add_picture(str(path), width=Inches(width_in))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(9)
    r.font.color.rgb = GREY
    cap.paragraph_format.space_after = Pt(10)


def style_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for cell in table.rows[0].cells:
        _set_cell_bg(cell, HEX_BLUE)
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.bold = True
                run.font.color.rgb = WHITE
                run.font.size = Pt(9.5)
    for r, row in enumerate(table.rows[1:], start=1):
        if r % 2 == 0:
            for cell in row.cells:
                _set_cell_bg(cell, HEX_ZEBRA)


def fill_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        cell.paragraphs[0].add_run(h).font.name = "Calibri"
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK
    style_table(table)
    return table


def add_footer_page_numbers(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Leistungspreis Industrie — Bayernwerk / REWAG · Seite ")
    r.font.size = Pt(8)
    r.font.color.rgb = GREY
    # PAGE-Feld
    run = p.add_run()
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    run._r.append(fld1); run._r.append(instr); run._r.append(fld2)
    run.font.size = Pt(8)
    run.font.color.rgb = GREY


# ===========================================================================
# Grafiken (matplotlib)
# ===========================================================================

def fig_cost_bars(path):
    fig, ax = plt.subplots(figsize=(5.2, 2.7))
    werte = [400 * 99.17, 500 * 99.17]
    bars = ax.bar(["400 kW Spitze", "500 kW Spitze"], werte,
                  color=[C_GREEN, C_RED], width=0.55)
    for b, v in zip(bars, werte):
        ax.text(b.get_x() + b.get_width() / 2, v + 800,
                f"{v:,.0f} €".replace(",", "."), ha="center",
                fontsize=10, fontweight="bold")
    ax.annotate("", xy=(1, werte[1]), xytext=(1, werte[0]),
                arrowprops=dict(arrowstyle="<->", color=C_GREY))
    ax.text(1.18, (werte[0] + werte[1]) / 2,
            "+9.917 €/Jahr", color=C_GREY, fontsize=9, va="center")
    ax.set_ylabel("Leistungskosten pro Jahr (€)")
    ax.set_title("Jahres-Leistungskosten — Bayernwerk MS, 99,17 €/(kW·a)",
                 fontsize=10, fontweight="bold")
    ax.set_ylim(0, 56000)
    ax.spines[["top", "right"]].set_visible(False)
    ax.grid(axis="y", alpha=0.25)
    fig.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def fig_hlzf_timeline(path):
    fig, ax = plt.subplots(figsize=(6.4, 1.5))
    ax.axvspan(0, 24, color=C_LIGHTGREEN)
    ax.axvspan(8.25, 18.25, color=C_RED, alpha=0.80)
    ax.text((8.25 + 18.25) / 2, 0.5, "Hochlast 08:15–18:15",
            ha="center", va="center", color="white", fontsize=9.5,
            fontweight="bold")
    ax.text(4.1, 0.5, "Nebenzeit", ha="center", va="center",
            color=C_GREEN, fontsize=9, fontweight="bold")
    ax.text(21.1, 0.5, "Nebenzeit", ha="center", va="center",
            color=C_GREEN, fontsize=9, fontweight="bold")
    ax.set_xlim(0, 24)
    ax.set_ylim(0, 1)
    ax.set_yticks([])
    ax.set_xticks(range(0, 25, 2))
    ax.set_xlabel("Uhrzeit (Werktag)")
    ax.set_title("Hochlastzeitfenster — Regensburg Netz, Mittelspannung, "
                 "Winter", fontsize=10, fontweight="bold")
    fig.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


def fig_solver_comparison(path):
    hours = list(range(24))
    # Normalfall: 240 kWh gleichmaessig -> 10 kW flach
    last_normal = [10] * 24
    # Atypisch: 0 im HLZF (8..17), Rest traegt die 240 kWh (14 h -> ~17 kW)
    last_atyp = []
    for h in hours:
        if 8 <= h < 18:
            last_atyp.append(0)
        else:
            last_atyp.append(240 / 14)

    fig, (ax1, ax2) = plt.subplots(2, 1, figsize=(6.4, 4.4), sharex=True)

    for ax in (ax1, ax2):
        ax.axvspan(8.0, 18.0, color=C_LIGHTRED, zorder=0)
        ax.set_xlim(-0.5, 23.5)
        ax.set_ylim(0, 22)
        ax.set_ylabel("Netzbezug (kW)")
        ax.spines[["top", "right"]].set_visible(False)
        ax.grid(axis="y", alpha=0.25)

    ax1.bar(hours, last_normal, color=C_PRIMARY, width=0.85)
    ax1.axhline(10, ls="--", color=C_RED, lw=1.2)
    ax1.text(0, 11.0, "Spitze 10 kW (rund um die Uhr)", color=C_RED,
             fontsize=8.5)
    ax1.set_title("Normalfall: Last über den Tag geglättet",
                  fontsize=10, fontweight="bold")

    ax2.bar(hours, last_atyp, color=C_PRIMARY, width=0.85)
    ax2.text(13.0, 2.0, "HLZF-Spitze = 0 kW", color=C_RED, ha="center",
             fontsize=9, fontweight="bold")
    ax2.text(21.3, 18.5, "außerhalb\nbeliebig", color=C_GREEN, ha="center",
             fontsize=8.5, fontweight="bold")
    ax2.set_title("Atypisch: Last aus dem HLZF (8–18 Uhr) verschoben",
                  fontsize=10, fontweight="bold")
    ax2.set_xticks(range(0, 24, 2))
    ax2.set_xlabel("Uhrzeit (Werktag)")

    legend = [Patch(facecolor=C_LIGHTRED, label="Hochlastzeitfenster")]
    ax1.legend(handles=legend, loc="upper right", fontsize=8, frameon=False)

    fig.tight_layout()
    fig.savefig(path, dpi=150, bbox_inches="tight")
    plt.close(fig)


# ===========================================================================
# Dokument
# ===========================================================================

def build_story(doc, charts: dict):
    add_banner(
        doc,
        "Leistungspreis für Industriekunden",
        "Normalfall und atypische Netznutzung — mit Beispielen, Grafiken "
        "und code-technischer Umsetzung · Bayernwerk Netz & REWAG/Regensburg, "
        "Stand 2026",
    )

    # ===================== Grundlagen =====================
    add_heading(doc, "Grundlagen (für beide Fälle gleich)")
    add_body(doc,
        "Privat- und Kleingewerbekunden (unter 100.000 kWh/Jahr) zahlen nur "
        "einen Arbeitspreis pro Kilowattstunde. Industriekunden mit "
        "registrierender Leistungsmessung (RLM, ab 100.000 kWh/Jahr) zahlen "
        "zusätzlich einen Leistungspreis: Euro pro Kilowatt und Jahr, bezogen "
        "auf die höchste 15-Minuten-Mittelleistung des Jahres."
    )
    add_body(doc,
        "Welcher Leistungspreis gilt, hängt von der Benutzungsdauer ab "
        "(Jahresverbrauch ÷ Jahresspitze). Ab 2.500 h/Jahr ist der "
        "Leistungspreis hoch und der Arbeitspreis niedrig, darunter "
        "umgekehrt. Das gilt in beiden Fällen — der Unterschied zwischen "
        "Normalfall und atypischer Netznutzung betrifft nur die Frage, "
        "WANN eine Spitze abrechnungsrelevant ist."
    )
    fill_table(doc,
        ["Benutzungsdauer", "Leistungspreis", "Arbeitspreis"],
        [
            ["≥ 2.500 h/Jahr", "hoch", "niedrig"],
            ["< 2.500 h/Jahr", "niedrig", "hoch"],
        ],
    )

    # ============================================================
    # TEIL A — NORMALFALL
    # ============================================================
    add_heading(doc, "Teil A — Normalfall: Jahresleistungspreis", size=16)

    add_sub(doc, "Konzept")
    add_body(doc,
        "Die Spitze zählt rund um die Uhr. Es ist völlig egal, ob die höchste "
        "Viertelstunde nachts, am Wochenende oder mittags auftritt — sie wird "
        "mit dem Leistungspreis multipliziert. Jede Spitze kostet, jederzeit."
    )

    add_sub(doc, "Reale Werte (netto)")
    add_body(doc, "Bayernwerk Netz, Mittelspannung 2026:",
             italic=True, space_after=3)
    fill_table(doc,
        ["Benutzungsdauer", "Leistungspreis €/(kW·a)", "Arbeitspreis ct/kWh"],
        [
            ["< 2.500 h/a", "15,86", "4,01"],
            ["≥ 2.500 h/a", "99,17", "0,68"],
        ],
    )
    add_body(doc, "Regensburg Netz (REWAG), Preisblatt 1 — Struktur:",
             italic=True, space_after=3)
    fill_table(doc,
        ["Ebene", "LP < 2500", "AP < 2500", "LP ≥ 2500", "AP ≥ 2500"],
        [
            ["Mittelspannung", "16,52", "4,90", "92,92", "1,84"],
            ["Umspannung MS/NS", "16,74", "5,61", "109,74", "1,89"],
            ["Niederspannung", "19,83", "5,91", "112,27", "2,18"],
        ],
    )

    add_sub(doc, "Beispiel")
    add_body(doc,
        "Ein Betrieb am Bayernwerk-Mittelspannungsnetz (≥ 2.500 h/a) zieht "
        "einmal im Jahr für 15 Minuten 500 kW. Bei 99,17 €/(kW·a) kostet "
        "allein diese Spitze rund 49.585 € pro Jahr. 100 kW weniger Spitze "
        "sparen fast 10.000 €:"
    )
    add_figure(doc, charts["cost"], 4.5,
               "Abb. 1: Jahres-Leistungskosten in Abhängigkeit von der "
               "Jahresspitze.")

    add_sub(doc, "Was bedeutet das für den Solver?")
    add_body(doc,
        "Der Solver muss die Spitze rund um die Uhr klein halten — Last über "
        "den ganzen Tag glätten, niemals eine hohe Einzelspitze zulassen. "
        "Im Test verteilt der Solver einen festen Tagesbedarf gleichmäßig "
        "über alle 24 Stunden (siehe obere Hälfte von Abb. 3).",
        space_after=8,
    )

    add_sub(doc, "Code-technische Umsetzung")
    add_body(doc,
        "Ein einziger Spitzen-Bucket über alle Zeitschritte. Im Solver eine "
        "Variable P_peak, an jeden Netzbezug gekoppelt:"
    )
    add_mono(doc, [
        "P_peak ≥ 0",
        "grid_buy[t] ≤ P_peak        # für ALLE Zeitschritte t",
        "Zielfunktion: + Leistungspreis · P_peak   (anteilig auf Horizont)",
    ])
    add_body(doc, "Konfiguration:", italic=True, space_after=3)
    add_mono(doc, [
        "cfg = {",
        "    'operator': 'bayernwerk',",
        "    'voltage_level': 'MS',",
        "    'benutzungsdauer': '>=2500',",
        "    'leistungspreis_mode': 'jahresleistungspreis',",
        "}",
    ])
    add_body(doc,
        "create_consumer_price.py baut daraus automatisch einen Bucket mit "
        "der Maske „alle Stunden“ und dem Leistungspreis aus der "
        "Tarif-Tabelle. demand_charge_constraints.py erzeugt die P_peak-"
        "Variable und den Zielfunktionsterm.", space_after=8,
    )

    # ============================================================
    # TEIL B — ATYPISCHER FALL
    # ============================================================
    add_heading(doc, "Teil B — Atypische Netznutzung (§ 19 StromNEV)", size=16)

    add_sub(doc, "Konzept")
    add_body(doc,
        "Nur die Spitze innerhalb der Hochlastzeitfenster (HLZF) zählt. Das "
        "sind die Zeiten, in denen das Netz stark belastet ist und die der "
        "Netzbetreiber jährlich vorgibt. Außerhalb dieser Fenster darf man "
        "beliebig hohe Spitzen fahren — sie sind nicht abrechnungsrelevant. "
        "Im Gegenzug erhält man ein individuell reduziertes Netzentgelt, das "
        "laut § 19 Abs. 2 Satz 1 StromNEV „nicht weniger als 20 Prozent des "
        "veröffentlichten Netzentgeltes“ betragen darf. Die Vereinbarung ist "
        "bei der Bundesnetzagentur anzuzeigen (Anzeigeverfahren nach "
        "BNetzA-Festlegung)."
    )

    add_sub(doc, "Reale Hochlastzeitfenster 2026")
    add_body(doc, "Regensburg Netz (REWAG), Werktage:",
             italic=True, space_after=3)
    fill_table(doc,
        ["Ebene", "Winter (Dez–Feb)", "Herbst (Sep–Nov)"],
        [
            ["Mittelspannung", "08:15 – 18:15", "09:30 – 12:30"],
            ["Umspannung MS/NS", "16:45 – 19:45", "17:00 – 19:00"],
            ["Niederspannung", "16:45 – 19:30", "17:00 – 19:00"],
        ],
    )
    add_figure(doc, charts["hlzf"], 6.2,
               "Abb. 2: Das HLZF teilt den Werktag in „teuer“ (Hochlast) und "
               "„gratis“ (Nebenzeit).")
    add_body(doc,
        "Bei Bayernwerk liegen die Fenster anders (z.B. NS Winter "
        "14:30–19:00). Die vollständigen Bayernwerk-HLZF stammen aus dem "
        "offiziellen Datenblatt 2026 und sind im Code als Datensatz "
        "hinterlegt.", space_after=8,
    )

    add_sub(doc, "Beispiel")
    add_body(doc,
        "Betrieb an Regensburg MS, gleiche 500-kW-Spitze, HLZF Winter "
        "08:15–18:15. Ob die Spitze etwas kostet, hängt jetzt vom Zeitpunkt ab:"
    )
    fill_table(doc,
        ["500 kW gezogen am …", "Normalfall", "Atypische Netznutzung"],
        [
            ["Dienstag 13:00 (im HLZF)", "teuer", "teuer"],
            ["Dienstag 22:00 (nach HLZF)", "teuer", "gratis"],
            ["Sonntag (kein Werktag)", "teuer", "gratis"],
        ],
    )

    add_sub(doc, "Was bedeutet das für den Solver?")
    add_body(doc,
        "Der Solver muss die Spitze nur in den HLZF drücken — dort möglichst "
        "auf null. Außerhalb darf er volle Leistung fahren. Praktisch: Last "
        "aus den teuren Fenstern heraus in die Nebenzeit (Abend, Nacht, "
        "Wochenende) verschieben."
    )
    add_figure(doc, charts["solver"], 6.2,
               "Abb. 3: Gleicher Tagesbedarf, zwei Modi. Oben Normalfall "
               "(geglättet), unten atypisch (HLZF leer).")

    add_sub(doc, "Code-technische Umsetzung")
    add_body(doc,
        "Ein einziger Spitzen-Bucket, dessen Maske nur die HLZF umfasst. Die "
        "Kopplung gilt dann nur für die Zeitschritte innerhalb der Fenster:"
    )
    add_mono(doc, [
        "P_peak ≥ 0",
        "grid_buy[t] ≤ P_peak        # NUR für t innerhalb der HLZF",
        "Zielfunktion: + Leistungspreis · P_peak",
    ])
    add_body(doc, "Konfiguration (nur der Modus ändert sich):",
             italic=True, space_after=3)
    add_mono(doc, [
        "cfg = {",
        "    'operator': 'regensburg',",
        "    'voltage_level': 'MS',",
        "    'benutzungsdauer': '>=2500',",
        "    'leistungspreis_mode': 'atypisch',",
        "}",
    ])
    add_body(doc,
        "create_consumer_price.py liest die HLZF des Netzbetreibers (Datensatz "
        "HLZF_2026, minutengenau inkl. Jahreszeit und Werktagslogik) und baut "
        "daraus die Bucket-Maske. Der Rest ist identisch zum Normalfall — nur "
        "die Maske ist enger.", space_after=8,
    )

    # ============================================================
    # GEMEINSAMER SCHLUSSTEIL
    # ============================================================
    add_heading(doc, "Gemeinsames: Dateien, Stolperfalle, Praxis", size=16)

    add_sub(doc, "Eine Datei, ein Aufruf")
    add_body(doc,
        "Der gesamte Code steckt in einer einzigen Datei (leistungspreis.py), "
        "die man in jedes PuLP-basierte Programm kopieren kann. Einzige "
        "Abhängigkeiten: numpy und pulp. Der Code selbst ist bewusst knapp "
        "kommentiert — was darin passiert, beschreibt dieser Abschnitt."
    )

    add_sub(doc, "Was in der Datei passiert")
    add_body(doc,
        "Die Datei gliedert sich in vier Bereiche, die in dieser Reihenfolge "
        "durchlaufen werden:"
    )
    fill_table(doc,
        ["Bereich", "Inhalt / Funktion"],
        [
            ["1. Tarifdaten",
             "TARIFFS: Leistungs- und Arbeitspreis je Netzbetreiber, "
             "Spannungsebene und Benutzungsdauer (echte Werte 2026)."],
            ["2. Hochlastzeitfenster",
             "HLZF_2026: die Zeitfenster je Betreiber/Ebene/Saison "
             "(minutengenau, nur Werktage) für die atypische Netznutzung."],
            ["3. build_price_and_demand()",
             "Rechnet ohne Solver: baut aus Spotpreis + Tarif den vollen "
             "kWh-Preis und legt die Spitzen-Definition (Bucket) an — je "
             "nach Modus über alle Stunden oder nur über die HLZF."],
            ["4. attach_leistungspreis()",
             "Die Solver-Anbindung: erzeugt die Variable P_peak, setzt die "
             "Nebenbedingungen grid_buy[t] ≤ P_peak ins Modell und gibt "
             "Preis-Zeitreihe und Leistungskosten-Term zurück."],
        ],
    )
    add_body(doc,
        "Zusätzlich: update_peak_so_far() liest nach dem Lösen die erreichte "
        "Spitze aus (für den rollierenden MPC), und classify_benutzungsdauer() "
        "ordnet bei Bedarf automatisch der richtigen Tarifvariante zu "
        "(über/unter 2.500 h/a).", space_after=8,
    )

    add_sub(doc, "Schritt 1–2: Importieren und einen Aufruf einfügen")
    add_mono(doc, [
        "from leistungspreis import attach_leistungspreis, update_peak_so_far",
        "",
        "TARIF = {",
        "    'operator': 'bayernwerk',          # oder 'regensburg'",
        "    'voltage_level': 'MS',",
        "    'benutzungsdauer': '>=2500',",
        "    'leistungspreis_mode': 'jahresleistungspreis',  # oder 'atypisch'",
        "}",
        "hook = attach_leistungspreis(model, grid_buy, spot_ct_kwh,",
        "                             timestamps, TARIF, dt_h=0.25)",
    ])

    add_sub(doc, "Schritt 3: Leistungspreis an die Zielfunktion anhängen")
    add_body(doc,
        "Deine bestehende Zielfunktion bleibt unverändert. Du hängst nur "
        "einen einzigen Term hinten an:"
    )
    add_mono(doc, [
        "# deine bestehende Zielfunktion bleibt wie sie ist …",
        "cost += hook.objective_term_ct      # … nur diese Zeile anfuegen",
        "model += cost",
    ])
    add_body(doc,
        "Zwei Hinweise: (1) Die Nebenbedingungen grid_buy[t] ≤ P_peak hat "
        "attach_leistungspreis() bereits in dein Modell eingefügt — das ist "
        "kein Eingriff in die Zielfunktion, aber nötig, damit P_peak die "
        "Spitze greift. (2) Der Term ist in ct (passend zu ct/kWh). Rechnet "
        "deine Zielfunktion in Euro, setze in der Datei oben "
        "_OBJECTIVE_UNIT_PER_EUR = 1.0."
    )
    add_body(doc, "Schritt 4 (nur bei rollierendem MPC, nach model.solve):",
             italic=True, space_after=3)
    add_mono(doc, [
        "peak_state = update_peak_so_far(hook)   # ins naechste Fenster geben",
    ])

    add_sub(doc, "Wie das Anhängen funktioniert")
    add_body(doc,
        "attach_leistungspreis() fasst die Zielfunktion selbst nicht an. Es "
        "baut nur den Term P_peak · Leistungspreis-Faktor und gibt ihn als "
        "objective_term_ct zurück. Das Anhängen ist normales Python: mit "
        "cost += … wird der Ausdruck zu deinen Kosten addiert, und PuLP setzt "
        "bei model += cost einen nackten Ausdruck als Zielfunktion (ein "
        "Vergleich wie ≤ wäre dagegen ein Constraint). Die Nebenbedingungen "
        "grid_buy[t] ≤ P_peak schreibt die Funktion direkt ins Modell — nur "
        "den Kostenbeitrag überlässt sie dir.", space_after=8,
    )

    add_sub(doc, "Inputs / Outputs auf einen Blick")
    fill_table(doc,
        ["Input (hineingeben)", "Output (hook.…)"],
        [
            ["model — dein LpProblem", "objective_term_ct — an Zielfunktion anhängen"],
            ["grid_buy — Netzbezug kW", "peak_vars — Spitze je Bucket (kW)"],
            ["spot_ct_kwh — Börsenpreis", "price_ct_kwh — voller kWh-Preis (optional)"],
            ["timestamps, TARIF, dt_h", "tariff / mode — Transparenz"],
        ],
    )

    add_sub(doc, "Stolperfalle: rollierender Horizont (beide Fälle)")
    add_body(doc,
        "Der Leistungspreis bezieht sich auf das ganze Jahr, die Optimierung "
        "läuft aber in kurzen Fenstern (z.B. 24–48 h). Damit ein Fenster die "
        "bereits erreichte Jahresspitze nicht vergisst, wird sie als "
        "Untergrenze (peak_so_far) in das nächste Fenster mitgegeben — sonst "
        "driftet der Peak nach oben.", space_after=8,
    )

    add_sub(doc, "Praxis-Hinweise")
    add_bullet(doc, "Netzentgelte und HLZF einmal jährlich aus dem aktuellen "
                    "Preisblatt aktualisieren (kein API-Abruf möglich).")
    add_bullet(doc, "Atypische Netznutzung ist bei der Bundesnetzagentur "
                    "anzuzeigen (§ 19 Abs. 2 StromNEV); Mindestentgelt 20 % "
                    "des veröffentlichten Netzentgelts.")
    add_bullet(doc, "Die Bayernwerk-HLZF für Mittelspannung (Winter/Herbst) "
                    "im Datensatz noch aus dem Originaldokument ergänzen.")

    add_sub(doc, "Faustregel")
    add_body(doc,
        "Spitze vermeiden, wenn sie Geld kostet — und nur dann. Im Normalfall "
        "kostet jede Spitze, also immer flach fahren. Mit atypischer "
        "Netznutzung nur in den Hochlastzeitfenstern flach fahren; außerhalb "
        "ist alles erlaubt."
    )

    # ===================== Quellen =====================
    add_heading(doc, "Quellen & Rechtsgrundlagen", size=16)
    add_body(doc,
        "Die regulatorischen Aussagen und Tarifwerte in diesem Dokument sind "
        "wie folgt belegt:"
    )
    fill_table(doc,
        ["Aussage", "Quelle / Fundstelle"],
        [
            ["Atypische Netznutzung — nur HLZF-Spitze zählt",
             "StromNEV § 19 Abs. 2 Satz 1"],
            ["Mindestentgelt „nicht weniger als 20 Prozent“",
             "StromNEV § 19 Abs. 2 Satz 1"],
            ["Anzeige bei der Regulierungsbehörde (Verfahren)",
             "StromNEV § 19 Abs. 2; BNetzA-Festlegung BK4-13-739"],
            ["Leistungspreis auf die 15-min-Jahreshöchstlast",
             "StromNEV § 17; Preisblätter der Netzbetreiber"],
            ["Bayernwerk MS: LP 99,17 / AP 0,68 (≥2500 h)",
             "Bayernwerk Netz, Elektronisches Preisblatt Strom 2026"],
            ["Regensburg: LP/AP-Struktur je Ebene",
             "Regensburg Netz (REWAG), Preisblatt 1"],
            ["Hochlastzeitfenster je Ebene & Saison",
             "HLZF-Datenblätter 2026 (Bayernwerk, Regensburg)"],
        ],
    )
    add_sub(doc, "Fundstellen (Stand 2026)")
    add_mono(doc, [
        "§ 19 StromNEV:",
        "  https://www.gesetze-im-internet.de/stromnev/__19.html",
        "Bundesnetzagentur — Individuelle Netzentgelte § 19 StromNEV:",
        "  https://www.bundesnetzagentur.de/DE/Beschlusskammern/BK04/",
        "  BK4_71_NetzE/BK4_71_Ind_NetzE_Strom/BK4_Ind_NetzEntg_Strom.html",
        "Bayernwerk Netz — Netzentgelte / HLZF 2026:",
        "  https://www.bayernwerk-netz.de/.../netzentgelte-strom.html",
        "Regensburg Netz (REWAG) — Netzentgelte / Preisblatt 1:",
        "  https://www.regensburg-netz.de/.../netzentgelte-umlagen",
    ])
    add_body(doc,
        "Hinweis: Netzentgelte und Hochlastzeitfenster ändern sich jährlich "
        "zum 1. Januar und werden bis 31. Oktober des Vorjahres "
        "veröffentlicht. Die im Code hinterlegten Werte sind entsprechend "
        "jährlich zu aktualisieren.", space_after=8,
    )


def main(out_path: Path) -> None:
    chart_dir = Path(tempfile.mkdtemp(prefix="lp_charts_"))
    charts = {
        "cost": chart_dir / "cost.png",
        "hlzf": chart_dir / "hlzf.png",
        "solver": chart_dir / "solver.png",
    }
    fig_cost_bars(charts["cost"])
    fig_hlzf_timeline(charts["hlzf"])
    fig_solver_comparison(charts["solver"])

    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    add_footer_page_numbers(doc)
    build_story(doc, charts)
    doc.save(str(out_path))
    print(f"DOCX erzeugt: {out_path}")

    # Charts aufraeumen (sind eingebettet)
    for p in charts.values():
        try:
            p.unlink()
        except OSError:
            pass
    try:
        chart_dir.rmdir()
    except OSError:
        pass


def _resolve_desktop() -> Path:
    home = Path.home()
    for c in (home / "OneDrive" / "Desktop",
              home / "OneDrive - Personal" / "Desktop",
              home / "Desktop"):
        if c.exists():
            return c
    raise FileNotFoundError("Kein Desktop-Ordner gefunden.")


if __name__ == "__main__":
    desktop = _resolve_desktop()
    target = (desktop / "leistungspreis_modul3" / "industrie"
              / "uebersicht_leistungspreis_industrie.docx")
    target.parent.mkdir(parents=True, exist_ok=True)
    try:
        main(target)
    except PermissionError:
        tmp = target.with_suffix(".new.docx")
        main(tmp)
        try:
            tmp.replace(target)
            print(f"Zieldatei ersetzt: {target}")
        except PermissionError:
            print(f"Hinweis: {target.name} ist gesperrt (Viewer offen). "
                  f"Neue Version liegt unter {tmp.name}.")
