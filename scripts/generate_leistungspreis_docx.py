"""Erzeugt eine praegnante Uebersichts-.docx zu Modul 3 und Leistungspreis."""

from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor


BLUE = RGBColor(0x0B, 0x3D, 0x91)
DARK = RGBColor(0x22, 0x22, 0x22)
GREY = RGBColor(0x55, 0x55, 0x55)


def _set_cell_bg(cell, hex_color: str) -> None:
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.makeelement(qn("w:shd"), {
        qn("w:val"): "clear", qn("w:color"): "auto", qn("w:fill"): hex_color,
    })
    tcpr.append(shd)


def add_heading(doc, text, size=15, color=BLUE, space_before=12, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_body(doc, text, size=10.5, italic=False, color=DARK, space_after=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.italic = italic
    run.font.color.rgb = color
    run.font.name = "Calibri"
    return p


def add_bullet(doc, text, size=10.5):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    run.font.color.rgb = DARK
    return p


def style_table(table, header_cells):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    # Kopfzeile faerben + fett
    for i, cell in enumerate(table.rows[0].cells):
        _set_cell_bg(cell, "0B3D91")
        for par in cell.paragraphs:
            for run in par.runs:
                run.font.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.font.size = Pt(9.5)
    # Zebra
    for r, row in enumerate(table.rows[1:], start=1):
        if r % 2 == 0:
            for cell in row.cells:
                _set_cell_bg(cell, "F2F6FB")


def fill_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.font.name = "Calibri"
    for row_data in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row_data):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(9.5)
            run.font.name = "Calibri"
            run.font.color.rgb = DARK
    style_table(table, headers)
    return table


def main(out_path: Path) -> None:
    doc = Document()

    # Standard-Schrift
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)

    # ---- Titel ----
    title = doc.add_paragraph()
    trun = title.add_run("Zeitvariable Netzentgelte & Leistungspreis")
    trun.bold = True
    trun.font.size = Pt(20)
    trun.font.color.rgb = BLUE
    title.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    srun = sub.add_run(
        "Modul 3 (§ 14a EnWG) und industrieller Leistungspreis — "
        "Konzept und Umsetzung in EMOS Light"
    )
    srun.italic = True
    srun.font.size = Pt(11)
    srun.font.color.rgb = GREY
    sub.paragraph_format.space_after = Pt(12)

    # ---- Worum es geht ----
    add_heading(doc, "Worum es geht")
    add_body(doc,
        "Ein einfacher Stromtarif rechnet einen festen Aufschlag pro "
        "Kilowattstunde. Zwei reale Preisbestandteile passen aber nicht in "
        "dieses Schema und brauchen je eine eigene Behandlung in der "
        "Optimierung. Beide setzen ein Preissignal, das Lastverschiebung "
        "belohnt — nur auf völlig unterschiedliche Weise."
    )

    # ---- Modul 1: Modul 3 ----
    add_heading(doc, "1. Modul 3 — zeitvariable Netzentgelte (Privat / § 14a EnWG)")
    add_body(doc,
        "Die Idee: Der Netzbetreiber teilt das Jahr in Zeitfenster mit drei "
        "Preisstufen ein — Hochlast (HT, teuer), Normallast (ST) und "
        "Niedriglast (NT, günstig). Wer seine steuerbaren Verbraucher "
        "(Wärmepumpe, Wallbox, Speicher) in die günstigen Fenster "
        "verschiebt, spart beim Netzentgelt. Die Fenster legt der "
        "Netzbetreiber pro Kalenderjahr fest und veröffentlicht sie im "
        "Voraus — sie hängen NICHT von der eigenen Last ab, sondern von "
        "der typischen Netzauslastung im Gebiet."
    )
    add_body(doc,
        "Wichtig für die Modellierung: Das ist weiterhin ein Preis in "
        "ct/kWh, nur eben zeitabhängig. Er lässt sich also direkt in die "
        "bestehende Preis-Zeitreihe einbauen — jeder Zeitschritt bekommt je "
        "nach Fenster seinen HT-, ST- oder NT-Satz."
    )
    add_body(doc, "Beispielwerte Bayernwerk 2026 (netto):", italic=True, space_after=3)
    fill_table(doc,
        ["Stufe", "Preis (ct/kWh)", "relativ zu ST"],
        [
            ["Hochlast (HT)", "9,03", "ca. +91 %"],
            ["Normaltarif (ST)", "4,72", "Referenz"],
            ["Niedriglast (NT)", "0,47", "ca. 10 % von ST"],
        ],
    )
    add_body(doc,
        "Beim Bayernwerk gelten die Zeitfenster 2026 nur im 2. und 3. Quartal "
        "(April bis September); in Herbst und Winter gilt durchgehend der "
        "Normaltarif.", space_after=8,
    )
    # (Hinweis: äöüß werden in diesem Dokument durchgehend verwendet.)

    # ---- Modul 2: Leistungspreis ----
    add_heading(doc, "2. Leistungspreis — Demand Charge (Industrie / RLM)")
    add_body(doc,
        "Die Idee: Industriebetriebe ab 100.000 kWh Jahresverbrauch zahlen "
        "zusätzlich zum Arbeitspreis einen Leistungspreis in Euro pro "
        "Kilowatt und Jahr — bezogen auf die höchste 15-Minuten-"
        "Mittelleistung des gesamten Jahres. Eine einzige Lastspitze im "
        "Januar verteuert also das ganze Jahr. Typisch sind 80 bis 180 "
        "€/(kW·a); bei 500 kW Spitze sind das schnell 40.000 bis 90.000 € "
        "pro Jahr. Deshalb lohnt sich Lastspitzenkappung (Peak Shaving) mit "
        "Speichern oder Lastmanagement."
    )
    add_body(doc,
        "Wichtig für die Modellierung: Das ist KEIN Preis pro kWh. Er hängt "
        "am Maximum aller Bezüge, nicht an einer einzelnen "
        "Kilowattstunde. Im MILP braucht das deshalb eine eigene Variable "
        "P_peak, die über Nebenbedingungen an jeden Netzbezug gekoppelt "
        "wird, plus einen eigenen Term in der Zielfunktion."
    )
    add_body(doc, "Kernmechanik im Solver:", italic=True, space_after=3)
    add_bullet(doc, "P_peak ist eine freie Variable (die später abgerechnete Spitze).")
    add_bullet(doc, "Für jeden relevanten Zeitschritt gilt: grid_buy[t] ≤ P_peak.")
    add_bullet(doc, "Zielfunktion: + Leistungspreis · P_peak (anteilig auf den Horizont).")
    add_body(doc,
        "Eine Variante ist die atypische Netznutzung (§ 19 StromNEV): Dann "
        "zählt nur die Spitze innerhalb der Hochlastzeitfenster — wer "
        "außerhalb seinen Peak hat, bekommt ein reduziertes Netzentgelt.",
        space_after=8,
    )

    # ---- Saisonale/zeitliche Differenzierung (Multi-Bucket) ----
    add_heading(doc, "3. Saisonale und zeitliche Differenzierung (Multi-Bucket)")
    add_body(doc,
        "Eine Lastspitze außerhalb der Hochlastzeitfenster ist "
        "wirtschaftlich anders zu bewerten als eine innerhalb. Deshalb "
        "wird der Leistungspreis nicht als eine einzige Spitze modelliert, "
        "sondern als mehrere Spitzen-\"Buckets\". Jeder Bucket hat ein "
        "eigenes Zeitfenster (Monat, Wochentag, Uhrzeit) und einen eigenen "
        "Preis in €/(kW·a)."
    )
    add_body(doc, "Im Solver bekommt jeder Bucket seine eigene Behandlung:",
             italic=True, space_after=3)
    add_bullet(doc, "eine eigene Variable P_peak je Bucket,")
    add_bullet(doc, "die Kopplung grid_buy[t] ≤ P_peak nur für die Zeitschritte "
                    "seiner Maske,")
    add_bullet(doc, "einen eigenen Term in der Zielfunktion mit dem Preis des Buckets.")
    add_body(doc,
        "So wird eine Spitze um 3 Uhr nachts automatisch \"billig\" und "
        "dieselbe Spitze um 18 Uhr im Winter \"teuer\" — ohne Extra-Regel, "
        "allein durch die unterschiedlichen Preise. Der Jahresleistungspreis "
        "ist der Spezialfall mit genau einem Bucket über alle Stunden; die "
        "atypische Netznutzung der Fall mit einem Bucket nur auf den "
        "Hochlastzeitfenstern."
    )
    add_body(doc, "Beispiele für die Bucket-Aufteilung:", italic=True, space_after=3)
    fill_table(doc,
        ["Tarifmodell", "Buckets"],
        [
            ["Jahresleistungspreis", "1 Bucket, Maske = alle Stunden"],
            ["Atypische Netznutzung (§ 19)",
             "1 Bucket nur auf HLZF; außerhalb gratis"],
            ["Haupt-/Nebenzeit getrennt",
             "2 disjunkte Buckets mit je eigenem Preis"],
        ],
        widths=[4000, 5360],
    )
    add_body(doc,
        "Wichtige Designentscheidung: Überlappen sich die Buckets (ein "
        "Schritt zählt zu mehreren → additive Bepreisung) oder sind sie "
        "disjunkt? Das hängt vom konkreten Preisblatt des Netzbetreibers "
        "ab und sollte am realen Tarif festgemacht werden. Saisonalität ist "
        "dabei kein Sonderfall, sondern steckt nur in der Definition der "
        "Zeitmaske (z.B. Monate November bis März).", space_after=8,
    )

    # ---- Gegenueberstellung ----
    add_heading(doc, "Die beiden im Vergleich")
    fill_table(doc,
        ["", "Modul 3 (zeitvariabel)", "Leistungspreis"],
        [
            ["Zielgruppe", "Privat / kleine Gewerbe", "Industrie / RLM"],
            ["Einheit", "ct/kWh (zeitabhängig)", "€/(kW·a) auf Jahres-Peak"],
            ["Bezug", "pro Zeitschritt", "Maximum über das Jahr"],
            ["Im MILP", "Teil der Preis-Zeitreihe", "eigene Variable + Term"],
            ["Hebel", "Last in günstige Fenster", "Spitzen kappen"],
        ],
    )

    # ---- Umsetzung im Code ----
    add_heading(doc, "Umsetzung im Code")
    add_body(doc,
        "Die Trennung der beiden Effekte spiegelt sich direkt in zwei "
        "Dateien wider:"
    )
    fill_table(doc,
        ["Datei", "Aufgabe"],
        [
            ["create_consumer_price.py",
             "Baut die zeitaufgelöste Preis-Zeitreihe (inkl. Modul-3-"
             "Zeitfenster) und liefert die Leistungspreis-Spezifikation "
             "(DemandChargeSpec mit den Spitzen-Buckets) — rechnet selbst "
             "nichts ab."],
            ["demand_charge_constraints.py",
             "Nimmt die Spezifikation und baut im Solver pro Bucket eine "
             "eigene P_peak-Variable, die Kopplungs-Constraints und den "
             "Zielfunktionsterm. Dieser Teil wandert in den Optimizer."],
        ],
        widths=[3000, 6360],
    )
    add_body(doc,
        "Die Zeitfenster werden regelbasiert hinterlegt (Monat, Wochentag, "
        "Uhrzeit) — kompakt abzutippen aus dem Preisblatt des "
        "Netzbetreibers. Derselbe Mechanismus bedient Modul 3 UND die "
        "Hochlastzeitfenster der atypischen Netznutzung.", space_after=8,
    )

    # ---- Stolperfalle ----
    add_heading(doc, "Wichtigste Stolperfalle: der rollierende Horizont")
    add_body(doc,
        "Der Leistungspreis bezieht sich auf das ganze Jahr, die "
        "Optimierung läuft aber in kurzen Fenstern (z.B. 24-48 h). Wenn "
        "jedes Fenster nur sich selbst optimiert, vergisst es die bereits "
        "erreichte Jahresspitze — und der Peak driftet langsam nach oben. "
        "Lösung: Die bisher erreichte Spitze (peak_so_far, je Bucket) wird "
        "als Untergrenze in jedes neue Fenster mitgegeben. So zahlt der "
        "Betrieb nur dann mehr, wenn das wirtschaftlich wirklich nötig ist.",
        space_after=8,
    )

    # ---- Fazit ----
    add_heading(doc, "Kurz gesagt")
    add_body(doc,
        "Modul 3 ist ein zeitabhängiger kWh-Preis und passt direkt in die "
        "bestehende Preislogik. Der Leistungspreis ist ein Maximum-Effekt "
        "und braucht eine eigene Variable im Solver. Beide lassen sich "
        "sauber getrennt umsetzen, ohne den restlichen Optimizer "
        "umzubauen — und beide kann man nicht per API abrufen, sondern muss "
        "sie einmal jährlich aus dem Preisblatt des Netzbetreibers pflegen."
    )

    doc.save(str(out_path))
    print(f"DOCX erzeugt: {out_path}")


def _resolve_desktop() -> Path:
    home = Path.home()
    for c in (home / "OneDrive" / "Desktop",
              home / "OneDrive - Personal" / "Desktop",
              home / "Desktop"):
        if c.exists():
            return c
    raise FileNotFoundError("Kein Desktop-Ordner gefunden.")


if __name__ == "__main__":
    desktop = _resolve_desktop()
    target = desktop / "leistungspreis_modul3" / "uebersicht_modul3_leistungspreis.docx"
    target.parent.mkdir(parents=True, exist_ok=True)
    main(target)
